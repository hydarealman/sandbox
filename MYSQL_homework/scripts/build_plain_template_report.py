from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = Path(r"C:\Users\dong\Desktop\数据库实验报告\生成的实验报告")
REPORT_PATH = REPORT_DIR / "数据库系统原理实验报告_软件工程3班_3242052051355_洪耀东.docx"

CLASS_NAME = "软件工程3班"
STUDENT_ID = "3242052051355"
STUDENT_NAME = "洪耀东"
DATE_TEXT = "2026年6月14日"


EXPERIMENTS = [
    {
        "title": "实验一 安装环境、简单编程与初识数据库",
        "purpose": [
            "掌握数据库运行环境和客户端工具的基本使用方法。",
            "掌握简单SQL程序设计，包括递归查询、数学计算和基础函数使用。",
            "创建订单管理数据库OrderDB，理解主键、外键和基础数据初始化过程。",
        ],
        "environment": [
            "Windows，本机MySQL80服务。",
            "MySQL 8.0.41，mysql命令行客户端。",
            "PowerShell、Python、PyMySQL、python-docx。",
        ],
        "content": [
            ("1.1  数据库环境检查", "在已配置好的MySQL环境中连接root用户，查看MySQL版本和当前数据库，确认数据库服务可用。"),
            ("1.2  SQL简单编程", "完成最大公约数、最小公倍数、水仙花数和分数数列前20项求和。"),
            ("1.3  OrderDB数据库初始化", "创建商品分类、商品、员工、客户、订单主表和订单明细表，并插入实验数据。"),
        ],
        "exercise": [
            ("1.1  源码文件", r"D:\sandbox\MYSQL_homework\sql\01_orderdb_setup_and_programming.sql"),
            ("1.2  运行结果", "脚本成功创建OrderDB，输出MySQL版本、表清单、程序题结果和基础数据量。"),
        ],
        "screens": ["screenshots/experiment_01.png"],
        "reflection": "本次实验把原指导书中的SQL Server环境内容迁移到MySQL实现，熟悉了MySQL数据类型、CHECK约束、递归CTE和外键约束的写法。",
    },
    {
        "title": "实验二 SQL查询",
        "purpose": [
            "掌握SELECT语句、条件查询、聚合函数和分组统计。",
            "掌握多表连接、子查询、EXISTS和复杂汇总查询。",
        ],
        "environment": [
            "Windows，本机MySQL80服务。",
            "MySQL 8.0.41，OrderDB数据库。",
            "mysql命令行客户端、PowerShell。",
        ],
        "content": [
            ("1.1  单表查询与聚合", "完成客户、员工、订单和订单明细相关的条件查询、排序、聚合和分组统计。"),
            ("1.2  多表查询", "完成内连接、左外连接、右外连接、UNION模拟完整外连接和子查询。"),
            ("1.3  复杂查询", "完成未订货客户、订购多种商品客户、销售员业绩、EXISTS相关查询等任务。"),
        ],
        "exercise": [
            ("1.1  源码文件", r"D:\sandbox\MYSQL_homework\sql\02_orderdb_queries.sql"),
            ("1.2  运行结果", "实验2-1、2-2、2-3查询均成功执行，结果日志保存在outputs目录中。"),
        ],
        "screens": ["screenshots/experiment_02.png"],
        "reflection": "复杂查询的关键是先明确查询粒度，再选择连接、分组或子查询。WHERE用于分组前过滤，HAVING用于聚合后过滤。",
    },
    {
        "title": "实验三 数据库与数据表定义、索引视图和数据更新",
        "purpose": [
            "掌握数据库、基本表、主键、外键和约束的定义方法。",
            "掌握索引、视图的创建、查询和删除。",
            "掌握INSERT、UPDATE、DELETE及其与约束之间的关系。",
        ],
        "environment": [
            "Windows，本机MySQL80服务。",
            "MySQL 8.0.41，BookDB数据库。",
            "mysql命令行客户端、PowerShell。",
        ],
        "content": [
            ("1.1  BookDB数据库定义", "创建图书分类、出版社、图书、读者、借阅五张表，并设置主键和外键。"),
            ("1.2  索引与视图", "创建出版社号、身份证号、工作单位等索引，创建BookView、BorrowView、ReaderView等视图。"),
            ("1.3  数据更新", "完成读者工作单位修改、图书价格调整、库存调整和借阅记录删除等操作。"),
        ],
        "exercise": [
            ("1.1  源码文件", r"D:\sandbox\MYSQL_homework\sql\03_bookdb_definition_indexes_views_updates.sql"),
            ("1.2  运行结果", "BookDB创建成功，索引、视图、查询和更新语句均执行成功。"),
        ],
        "screens": ["screenshots/experiment_03.png"],
        "reflection": "视图能简化复杂连接查询，索引能提高检索效率。进行数据更新前应先分析外键和业务约束，避免破坏数据一致性。",
    },
    {
        "title": "实验四 游标、存储过程与触发器",
        "purpose": [
            "掌握MySQL游标和存储过程的创建、执行和调用方法。",
            "掌握触发器在安全控制、数据校验和库存维护中的应用。",
        ],
        "environment": [
            "Windows，本机MySQL80服务。",
            "MySQL 8.0.41，OrderDB数据库。",
            "mysql命令行客户端、PowerShell。",
        ],
        "content": [
            ("1.1  游标与存储过程", "使用游标统计订单金额，编写员工编号生成、大客户热销商品、客户销售汇总和发票输出存储过程。"),
            ("1.2  触发器", "编写客户删除控制、薪水修改审计、员工日期校验和订单明细库存维护触发器。"),
        ],
        "exercise": [
            ("1.1  源码文件", r"D:\sandbox\MYSQL_homework\sql\04_orderdb_procedures_triggers.sql"),
            ("1.2  运行结果", "存储过程调用成功，触发器生成薪水审计记录，并在插入订单明细后自动更新库存和订单金额。"),
        ],
        "screens": ["screenshots/experiment_04.png"],
        "reflection": "游标适合逐行处理过程化任务，触发器适合把审计、校验和库存维护封装到数据库内部，减少应用层遗漏。",
    },
    {
        "title": "实验五 安全性定义与检查",
        "purpose": [
            "掌握数据库用户、角色和权限的创建、授权、回收和检查。",
            "理解MySQL权限模型与SQL Server用户模型之间的差异。",
        ],
        "environment": [
            "Windows，本机MySQL80服务。",
            "MySQL 8.0.41，OrderDB数据库。",
            "mysql命令行客户端、PowerShell。",
        ],
        "content": [
            ("1.1  用户与角色创建", "创建user01、user02、user03、user05、user06、user07等用户，并创建r3、r5、r6等角色。"),
            ("1.2  授权与回收", "完成表级权限、列级权限、角色权限、建表和建视图权限的授予与回收。"),
            ("1.3  权限检查", "使用SHOW GRANTS检查用户权限，创建测试表和视图验证权限对象。"),
        ],
        "exercise": [
            ("1.1  源码文件", r"D:\sandbox\MYSQL_homework\sql\05_security_mysql.sql"),
            ("1.2  运行结果", "用户、角色和授权语句执行成功，SHOW GRANTS输出显示权限分配结果。"),
        ],
        "screens": ["screenshots/experiment_05.png"],
        "reflection": "数据库安全控制不仅是能否登录，还包括对象级、列级和角色继承等细粒度权限管理。",
    },
    {
        "title": "实验六 完整性定义与检查",
        "purpose": [
            "掌握实体完整性、参照完整性和用户自定义完整性的定义方法。",
            "通过插入、删除、更新操作观察完整性约束的检查效果。",
        ],
        "environment": [
            "Windows，本机MySQL80服务。",
            "MySQL 8.0.41，BookDB数据库。",
            "mysql命令行客户端、PowerShell。",
        ],
        "content": [
            ("1.1  完整性约束定义", "重新创建BookDB五张表，设置主键、外键、唯一约束、CHECK约束和默认值。"),
            ("1.2  数据插入", "插入5条图书分类、5条出版社、20条图书、10条读者和50条借阅记录。"),
            ("1.3  完整性检查", "测试重复主键、重复身份证、外键限制、日期限制、最大借书数和价格范围约束。"),
        ],
        "exercise": [
            ("1.1  源码文件", r"D:\sandbox\MYSQL_homework\sql\06_bookdb_integrity.sql"),
            ("1.2  运行结果", "合法操作执行成功，违反主键、外键或CHECK约束的操作被MySQL完整性机制拦截。"),
        ],
        "screens": ["screenshots/experiment_06.png"],
        "reflection": "完整性约束能把业务规则固定在数据库层，减少应用程序遗漏校验造成的数据错误。",
    },
    {
        "title": "实验七 执行计划",
        "purpose": [
            "掌握使用EXPLAIN查看SQL查询执行计划的方法。",
            "理解访问类型、候选索引、实际索引、扫描行数和Extra信息对SQL优化的意义。",
        ],
        "environment": [
            "Windows，本机MySQL80服务。",
            "MySQL 8.0.41，OrderDB数据库。",
            "mysql命令行客户端、PowerShell。",
        ],
        "content": [
            ("1.1  索引准备", "为商品名称、订单明细商品号、订单号、订单主表员工号和客户号建立索引。"),
            ("1.2  执行计划分析", "查看“酷睿四核i5-6500”销售情况查询和每个员工销售记录查询的执行计划。"),
        ],
        "exercise": [
            ("1.1  源码文件", r"D:\sandbox\MYSQL_homework\sql\07_execution_plan.sql"),
            ("1.2  运行结果", "EXPLAIN结果显示查询使用了相关索引，并输出对应销售查询结果。"),
        ],
        "screens": ["screenshots/experiment_07.png"],
        "reflection": "SQL优化应先保证结果正确，再利用EXPLAIN观察是否走索引、是否产生临时表或文件排序，并据此调整SQL或索引。",
    },
    {
        "title": "实验八 事务处理",
        "purpose": [
            "理解事务提交、回滚和隔离级别的作用。",
            "掌握START TRANSACTION、COMMIT以及多连接并发验证方法。",
        ],
        "environment": [
            "Windows，本机MySQL80服务。",
            "MySQL 8.0.41，OrderDB数据库。",
            "mysql命令行客户端、PowerShell、Python、PyMySQL。",
        ],
        "content": [
            ("1.1  事务提交", "定义事务注册新客户，并定义事务登记2022年8月16日的新订单，新订单享受2%优惠。"),
            ("1.2  事务更新", "定义事务处理业务员E2020003离职，解除订单引用后删除员工信息。"),
            ("1.3  隔离级别", "通过两个数据库连接模拟T1查询平均薪水、T2插入新员工并提交，比较READ COMMITTED和REPEATABLE READ效果。"),
        ],
        "exercise": [
            ("1.1  源码文件", r"D:\sandbox\MYSQL_homework\sql\08_transactions.sql；D:\sandbox\MYSQL_homework\scripts\run_transaction_isolation_demo.py"),
            ("1.2  运行结果", "新客户、新订单和离职处理事务执行成功；并发演示显示不同隔离级别下两次读取结果不同。"),
        ],
        "screens": ["screenshots/experiment_08.png", "screenshots/experiment_09.png"],
        "reflection": "事务隔离级别会影响同一事务内多次查询能否看到其他事务提交的数据，实际应用中需要根据一致性要求选择合适级别。",
    },
]


def apply_run_font(run, size=11, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_text(doc, text="", size=11, bold=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    apply_run_font(run, size=size, bold=bold)
    return p


def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    apply_run_font(run, size=12, bold=True)


def add_numbered_lines(doc, lines):
    for index, line in enumerate(lines, 1):
        add_text(doc, f"{index}. {line}", size=11)


def add_label_value(doc, label, value):
    add_text(doc, label, size=11, bold=True)
    add_text(doc, value, size=11)


def add_header(doc, exp_title):
    add_text(doc, "桂 林 理 工 大 学", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "实  验  报  告", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, f"班级  {CLASS_NAME}          学号  {STUDENT_ID}          姓名  {STUDENT_NAME}", size=11)
    add_text(doc, f"实验名称  {exp_title}          日期  {DATE_TEXT}", size=11)


def build_report():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)

    for exp_index, exp in enumerate(EXPERIMENTS, 1):
        if exp_index > 1:
            doc.add_page_break()

        add_header(doc, exp["title"])

        add_section_title(doc, "一、实验目的：")
        add_numbered_lines(doc, exp["purpose"])

        add_section_title(doc, "二、实验环境：")
        add_numbered_lines(doc, exp["environment"])

        add_section_title(doc, "三、实验内容：")
        for title, detail in exp["content"]:
            add_label_value(doc, title, detail)

        add_section_title(doc, "四、实验练习")
        for title, detail in exp["exercise"]:
            add_label_value(doc, title, detail)

        for idx, screen in enumerate(exp["screens"], 1):
            add_text(doc, f"运行截图 {exp_index}.{idx}：", size=11, bold=True)
            pic = doc.add_picture(str(ROOT / screen), width=Inches(6.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        add_section_title(doc, "五、实验心得")
        add_text(doc, exp["reflection"], size=11)

    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_PATH)
    print(REPORT_PATH)


if __name__ == "__main__":
    build_report()
