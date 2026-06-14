from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_plain_template_report import (
    CLASS_NAME,
    DATE_TEXT,
    EXPERIMENTS,
    ROOT,
    STUDENT_ID,
    STUDENT_NAME,
    add_header,
    add_label_value,
    add_numbered_lines,
    add_section_title,
    add_text,
)


REPORT_DIR = Path(r"C:\Users\dong\Desktop\数据库实验报告\生成的实验报告\8个单独实验报告")


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)
    return doc


def build_one(exp_index, exp):
    doc = setup_document()
    add_header(doc, exp["title"])

    add_section_title(doc, "一、实验目的：")
    add_numbered_lines(doc, exp["purpose"])

    add_section_title(doc, "二、实验环境：")
    add_numbered_lines(doc, exp["environment"])

    add_section_title(doc, "三、实验内容：")
    for title, detail in exp["content"]:
        add_label_value(doc, title, detail)

    add_section_title(doc, "四、实验练习")
    for title, detail in exp["exercise"]:
        add_label_value(doc, title, detail)

    for idx, screen in enumerate(exp["screens"], 1):
        add_text(doc, f"运行截图 {exp_index}.{idx}：", size=11, bold=True)
        doc.add_picture(str(ROOT / screen), width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_section_title(doc, "五、实验心得")
    add_text(doc, exp["reflection"], size=11)
    return doc


def main():
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    generated = []
    for index, exp in enumerate(EXPERIMENTS, 1):
        doc = build_one(index, exp)
        name = f"实验{index}_{exp['title'].replace(' ', '_')}_{CLASS_NAME}_{STUDENT_ID}_{STUDENT_NAME}.docx"
        # Avoid filesystem surprises from punctuation in experiment titles.
        for bad in "\\/:*?\"<>|":
            name = name.replace(bad, "_")
        path = REPORT_DIR / name
        doc.save(path)
        generated.append(path)

    manifest = REPORT_DIR / "文件清单.txt"
    manifest.write_text("\n".join(str(path) for path in generated), encoding="utf-8")
    for path in generated:
        print(path)


if __name__ == "__main__":
    main()
