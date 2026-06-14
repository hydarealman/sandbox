from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DESKTOP_REPORT_DIR = Path(r"C:\Users\dong\Desktop\数据库实验报告\生成的实验报告")
REPORT_PATH = DESKTOP_REPORT_DIR / "数据库系统原理实验报告_软件工程3班_3242052051355_洪耀东.docx"


STUDENT = {
    "class": "软件工程3班",
    "student_id": "3242052051355",
    "name": "洪耀东",
    "date": "2026年6月14日",
}


EXPERIMENTS = [
    {
        "name": "实验一 安装环境、简单编程与初识数据库",
        "purpose": [
            "掌握数据库运行环境和客户端工具的基本使用方法。",
            "掌握变量、递推查询和简单程序逻辑的实现方法。",
            "创建订单管理数据库OrderDB，理解主键、外键和基础数据插入过程。",
        ],
        "content": [
            "在已配置好的MySQL 8.0环境中验证服务和客户端可用性。",
            "编写求最大公约数、最小公倍数、水仙花数和分数数列求和的SQL程序。",
            "建立ProductClass、Product、Employee、Customer、OrderMaster、OrderDetail六张表并插入样例数据。",
        ],
        "source": "sql/01_orderdb_setup_and_programming.sql",
        "screens": ["screenshots/experiment_01.png"],
        "result": "脚本成功创建OrderDB，完成基础数据初始化，并输出MySQL版本、表清单和程序题结果。",
        "reflection": "本次实验把原指导书中的SQL Server环境要求迁移到MySQL实现，重点体会了数据类型、CHECK约束和递归CTE的写法差异。",
    },
    {
        "name": "实验二 SQL查询",
        "purpose": [
            "熟练掌握SELECT语句、聚合函数、分组统计和排序。",
            "掌握多表连接、子查询、EXISTS和复杂汇总查询。",
        ],
        "content": [
            "完成客户、员工、订单和订单明细的单表查询。",
            "完成内连接、左右外连接、UNION模拟完整外连接和子查询。",
            "完成复杂查询，包括未订货客户、至少订购多种商品客户和销售业绩统计。",
        ],
        "source": "sql/02_orderdb_queries.sql",
        "screens": ["screenshots/experiment_02.png"],
        "result": "脚本完成实验2-1、2-2、2-3全部查询，并给出GROUP BY、WHERE、HAVING、EXISTS等问题简答。",
        "reflection": "复杂查询的关键是先确定数据粒度，再决定使用连接、分组还是相关子查询；WHERE用于分组前过滤，HAVING用于聚合后过滤。",
    },
    {
        "name": "实验三 数据库与数据表定义、索引视图和数据更新",
        "purpose": [
            "掌握数据库、基本表、主键、外键和约束的定义方法。",
            "掌握索引和视图的创建、查询与删除。",
            "掌握INSERT、UPDATE、DELETE及其与约束之间的关系。",
        ],
        "content": [
            "创建BookDB及图书分类、出版社、图书、读者、借阅五张表。",
            "建立出版社号、身份证号、工作单位和最大借书数量相关索引。",
            "创建BookView、BorrowView、ReaderView并执行查询。",
            "完成工作单位修改、价格调整、库存调整、借阅记录删除和从未借书读者删除。",
        ],
        "source": "sql/03_bookdb_definition_indexes_views_updates.sql",
        "screens": ["screenshots/experiment_03.png"],
        "result": "BookDB创建成功，视图查询和数据更新均成功执行，日志显示更新后的读者数、借阅数和经济类图书数量。",
        "reflection": "视图能简化复杂连接查询，索引能为查询优化提供基础；更新操作必须先考虑外键和业务约束，避免破坏数据一致性。",
    },
    {
        "name": "实验四 游标、存储过程与触发器",
        "purpose": [
            "掌握MySQL游标和存储过程的定义、调用方法。",
            "掌握触发器在安全审计、数据校验和库存维护中的应用。",
        ],
        "content": [
            "使用游标重新统计OrderMaster中的订单金额。",
            "编写自动生成员工编号、查询大客户热销商品、客户销售汇总和发票信息的存储过程。",
            "编写客户删除控制、薪水修改审计、日期合法性检查和订单明细库存维护触发器。",
        ],
        "source": "sql/04_orderdb_procedures_triggers.sql",
        "screens": ["screenshots/experiment_04.png"],
        "result": "存储过程调用成功，触发器记录薪水审计日志，并在插入订单明细后自动更新库存和订单金额。",
        "reflection": "游标适合逐行处理格式化输出或过程控制任务，触发器适合把一致性维护逻辑封装到数据库内部。",
    },
    {
        "name": "实验五 安全性定义与检查",
        "purpose": [
            "掌握数据库用户、角色和权限的创建、授权、回收和检查。",
            "理解MySQL权限模型与SQL Server登录用户模型之间的差异。",
        ],
        "content": [
            "创建user01、user02、user03、user05、user06、user07等用户。",
            "创建r3、r5、r6角色并分配表权限、列权限和建表建视图权限。",
            "使用SHOW GRANTS检查授权结果，并创建测试表和测试视图。",
        ],
        "source": "sql/05_security_mysql.sql",
        "screens": ["screenshots/experiment_05.png"],
        "result": "用户、角色和权限成功创建，SHOW GRANTS输出展示了各用户权限和角色成员关系。",
        "reflection": "权限实验说明数据库安全控制不仅包含能否登录，还包含对象级、列级以及角色继承关系的细粒度控制。",
    },
    {
        "name": "实验六 完整性定义与检查",
        "purpose": [
            "掌握实体完整性、参照完整性和用户自定义完整性的定义方式。",
            "通过实际插入、删除、更新操作观察完整性约束的拦截效果。",
        ],
        "content": [
            "重新创建BookDB五张表，增加主键、外键、唯一约束、CHECK约束和默认值。",
            "插入5条图书分类、5条出版社、20条图书、10条读者和50条借阅记录。",
            "用存储过程批量测试重复主键、重复身份证、外键限制、日期限制、最大借书数和价格范围约束。",
        ],
        "source": "sql/06_bookdb_integrity.sql",
        "screens": ["screenshots/experiment_06.png"],
        "result": "完整性测试日志显示：合法操作执行成功，违反主键、外键或CHECK约束的操作被MySQL完整性机制拦截。",
        "reflection": "完整性约束能把业务规则前移到数据库层，减少应用程序遗漏校验带来的脏数据风险。",
    },
    {
        "name": "实验七 执行计划",
        "purpose": [
            "掌握使用EXPLAIN查看SQL执行计划的方法。",
            "理解访问类型、候选索引、实际索引、扫描行数和Extra信息对优化的意义。",
        ],
        "content": [
            "为商品名称、订单明细商品号、订单号、订单主表员工号和客户号建立索引。",
            "查看“酷睿四核i5-6500”销售情况查询的执行计划。",
            "查看每个员工销售记录查询的执行计划并输出查询结果。",
        ],
        "source": "sql/07_execution_plan.sql",
        "screens": ["screenshots/experiment_07.png"],
        "result": "EXPLAIN结果显示相关查询使用了商品名称、订单明细商品号、订单号等索引，查询结果与业务数据一致。",
        "reflection": "优化SQL时应先写出正确查询，再结合EXPLAIN判断是否走索引、是否产生临时表或文件排序。",
    },
    {
        "name": "实验八 事务处理",
        "purpose": [
            "理解事务提交、回滚和隔离级别的作用。",
            "掌握START TRANSACTION、COMMIT以及多连接并发验证方法。",
        ],
        "content": [
            "定义事务注册新客户，并定义事务登记2022年8月16日新订单，订单享受2%优惠。",
            "定义事务处理业务员E2020003离职，解除订单引用后删除员工信息。",
            "通过Python脚本创建两个数据库连接，模拟T1查询平均薪水、T2插入新员工并提交，比较READ COMMITTED和REPEATABLE READ效果。",
        ],
        "source": "sql/08_transactions.sql；scripts/run_transaction_isolation_demo.py",
        "screens": ["screenshots/experiment_08.png", "screenshots/experiment_09.png"],
        "result": "事务脚本成功完成新客户、新订单和离职业务员处理；并发演示显示READ COMMITTED第二次读到新数据，REPEATABLE READ两次读取一致。",
        "reflection": "事务隔离级别会直接影响同一事务内多次查询能否看到其他事务提交的数据，实际系统中需要在一致性和并发性能之间取舍。",
    },
]

OUTPUT_FILES = [
    "outputs/01_orderdb_setup_and_programming.txt",
    "outputs/02_orderdb_queries.txt",
    "outputs/03_bookdb_definition_indexes_views_updates.txt",
    "outputs/04_orderdb_procedures_triggers.txt",
    "outputs/05_security_mysql.txt",
    "outputs/06_bookdb_integrity.txt",
    "outputs/07_execution_plan.txt",
    "outputs/08_transactions.txt",
    "outputs/08_isolation_demo.txt",
]


def set_cell_text(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "B7C3D0")
        borders.append(tag)
    tbl_pr.append(borders)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)


def add_label_paragraph(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label)
    r.bold = True
    p.add_run(text)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    run = title.add_run("桂 林 理 工 大 学")
    run.bold = True
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("实  验  报  告")
    run.bold = True
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(20)

    meta = doc.add_table(rows=2, cols=3)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(meta)
    entries = [
        ("班级", STUDENT["class"]),
        ("学号", STUDENT["student_id"]),
        ("姓名", STUDENT["name"]),
        ("实验名称", "数据库系统原理综合实验"),
        ("日期", STUDENT["date"]),
        ("环境", "Windows + MySQL 8.0.41"),
    ]
    for cell, (label, value) in zip(meta._cells, entries):
        set_cell_text(cell, f"{label}：{value}", bold=True)

    doc.add_paragraph()
    add_label_paragraph(doc, "说明：", "原实验指导书以SQL Server 2019为环境，本报告按用户已配置的MySQL 8.0环境完成等价实现。源码统一保存于 D:\\sandbox\\MYSQL_homework。")

    doc.add_heading("源码清单", level=1)
    for index, exp in enumerate(EXPERIMENTS, 1):
        label = f"实验{index if index <= 8 else '8补充'}"
        add_bullet(doc, f"{label}：{exp['source']}；输出：{OUTPUT_FILES[index - 1]}")

    for idx, exp in enumerate(EXPERIMENTS, 1):
        doc.add_page_break()
        doc.add_heading(exp["name"], level=1)

        meta = doc.add_table(rows=2, cols=3)
        meta.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(meta)
        entries = [
            ("班级", STUDENT["class"]),
            ("学号", STUDENT["student_id"]),
            ("姓名", STUDENT["name"]),
            ("实验名称", exp["name"]),
            ("日期", STUDENT["date"]),
            ("源码", exp["source"]),
        ]
        for cell, (label, value) in zip(meta._cells, entries):
            set_cell_text(cell, f"{label}：{value}", bold=True if label in {"班级", "学号", "姓名"} else False)

        doc.add_heading("一、实验目的", level=2)
        for item in exp["purpose"]:
            add_bullet(doc, item)

        doc.add_heading("二、实验环境", level=2)
        add_bullet(doc, "操作系统：Windows，本机MySQL服务 MySQL80。")
        add_bullet(doc, "数据库：MySQL 8.0.41，字符集utf8mb4。")
        add_bullet(doc, "工具：mysql命令行客户端、PowerShell、Python、PyMySQL、python-docx。")

        doc.add_heading("三、实验内容", level=2)
        for item in exp["content"]:
            add_bullet(doc, item)

        doc.add_heading("四、实验练习与运行结果", level=2)
        add_label_paragraph(doc, "源码位置：", f"D:\\sandbox\\MYSQL_homework\\{exp['source']}")
        add_label_paragraph(doc, "运行结果：", exp["result"])
        for screen in exp["screens"]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(f"图 {idx}-{exp['screens'].index(screen) + 1} 源码与运行输出截图")
            run.bold = True
            doc.add_picture(str(ROOT / screen), width=Inches(6.3))

        doc.add_heading("五、实验心得", level=2)
        doc.add_paragraph(exp["reflection"])

    DESKTOP_REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_PATH)
    return REPORT_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
