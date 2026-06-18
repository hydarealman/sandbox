#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 lab5 实验报告 Word 文档 (含终端截图)
实验5: Java输入输出流与多线程编程
"""

import os
import subprocess
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont

LAB5_DIR = r"d:\sandbox\JAVALAB\lab5"
SCREENSHOTS_DIR = os.path.join(LAB5_DIR, "screenshots")

def create_terminal_screenshot(text, output_path, title="", width=900, max_height=1200):
    """使用PIL生成终端样式的截图"""
    os.makedirs(SCREENSHOTS_DIR, exist_ok=True)

    # 如果没有文本，使用占位符
    if not text:
        text = "(no output)"

    lines = text.split('\n')
    # 让所有行都能显示，自动扩展高度
    line_height = 18
    padding = 20
    title_height = 30 if title else 0
    # 计算需要的高度（最小300，最大max_height）
    needed_height = min(max(len(lines) * line_height + padding * 2 + title_height + 10, 200), max_height)
    # 如果行数超过了max_height容纳范围，截断文本
    max_lines = (needed_height - padding * 2 - title_height) // line_height
    if len(lines) > max_lines:
        lines = lines[:max_lines]
        lines.append("... (output truncated)")

    img_height = padding * 2 + title_height + len(lines) * line_height + 10
    img_height = max(200, min(img_height, max_height))

    img = Image.new('RGB', (width, img_height), color=(30, 30, 30))  # 终端深色背景
    draw = ImageDraw.Draw(img)

    # 尝试加载中文字体
    font_paths = [
        "C:\\Windows\\Fonts\\msyh.ttc",       # 微软雅黑
        "C:\\Windows\\Fonts\\simsun.ttc",      # 宋体
        "C:\\Windows\\Fonts\\consola.ttf",     # Consolas
        "C:\\Windows\\Fonts\\simhei.ttf",      # 黑体
    ]
    font = None
    for fp in font_paths:
        if os.path.exists(fp):
            try:
                font = ImageFont.truetype(fp, 13)
                break
            except:
                continue
    if font is None:
        font = ImageFont.load_default()

    # 标题栏
    if title:
        draw.rectangle([(0, 0), (width, title_height)], fill=(60, 60, 60))
        draw.text((padding, 5), title, fill=(255, 255, 255), font=font)

    # 绘制终端文本
    y = padding + title_height
    for line in lines:
        if y > img_height - padding:
            break
        # 跳过空行或处理
        if line.strip() == "":
            draw.text((padding, y), " ", fill=(200, 200, 200), font=font)
        elif "==========" in line or "---" in line:
            draw.text((padding, y), line, fill=(100, 255, 100), font=font)
        elif "正确" in line or "恭喜" in line or "对了" in line:
            draw.text((padding, y), line, fill=(100, 255, 100), font=font)
        elif "错误" in line or "不正确" in line or "错了" in line:
            draw.text((padding, y), line, fill=(255, 100, 100), font=font)
        elif "总成绩" in line or "分数" in line:
            draw.text((padding, y), line, fill=(255, 255, 100), font=font)
        elif "已答" in line or "正确率" in line:
            draw.text((padding, y), line, fill=(100, 200, 255), font=font)
        else:
            draw.text((padding, y), line, fill=(200, 200, 200), font=font)
        y += line_height

    img.save(output_path)
    return output_path

def add_formatted_paragraph(doc, text, font_name='宋体', font_size=12, bold=False, alignment=None, color=None):
    """添加格式化段落"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_code_block(doc, code_text, title=None):
    """添加代码块"""
    if title:
        add_formatted_paragraph(doc, title, font_name='黑体', font_size=11, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 80, 0)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    shading.set(qn('w:val'), 'clear')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    return p

def add_screenshot(doc, image_path, caption="", width_inches=5.8):
    """添加截图到Word文档"""
    if caption:
        add_formatted_paragraph(doc, caption, font_name='黑体', font_size=11, bold=True,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(width_inches))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_formatted_paragraph(doc, f'[截图未找到: {image_path}]', font_size=10, color=(255,0,0))
    doc.add_paragraph()

def run_java(class_name, input_text=None, timeout=10):
    """运行 Java 程序并捕获输出"""
    try:
        cmd = ['java', f'JAVALAB.lab5.{class_name}']
        result = subprocess.run(
            cmd, cwd=r'd:\sandbox', input=input_text,
            capture_output=True, text=True, timeout=timeout,
            encoding='utf-8', errors='replace'
        )
        output = result.stdout
        if result.stderr:
            output += "\n" + result.stderr
        return output.strip()
    except subprocess.TimeoutExpired:
        return "[程序超时]"
    except Exception as e:
        return f"[运行错误: {e}]"

def run_and_screenshot(class_name, screenshot_name, title, input_text=None, timeout=10):
    """运行Java程序，捕获输出，生成截图"""
    output = run_java(class_name, input_text, timeout)
    if output:
        img_path = create_terminal_screenshot(
            output,
            os.path.join(SCREENSHOTS_DIR, f"{screenshot_name}.png"),
            title=title
        )
        return output, img_path
    return output, None

def main():
    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ============ 页码边距 ============
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ============ 标题 ============
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('桂 林 理 工 大 学')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(22)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('实  验  报  告')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(22)
    run.bold = True
    doc.add_paragraph()

    # ============ 信息表 ============
    table = doc.add_table(rows=2, cols=8, style='Table Grid')
    table.autofit = True

    cells_r1 = table.rows[0].cells
    cells_r1[0].text = '班级'
    cells_r1[1].text = '软工3班'
    cells_r1[2].text = '学号'
    cells_r1[3].text = '3242052051355'
    cells_r1[4].text = '姓名'
    cells_r1[5].text = '洪耀东'
    cells_r1[6].text = '同组实验者'
    cells_r1[7].text = ''

    cells_r2 = table.rows[1].cells
    cells_r2[0].text = '实验名称'
    cells_r2[1].text = '面向对象程序设计  实验5 Java输入输出流与多线程编程'
    cells_r2[2].text = ''
    cells_r2[3].text = ''
    cells_r2[4].text = ''
    cells_r2[5].text = '日期'
    cells_r2[6].text = '2026年 04月 18 日'
    cells_r2[7].text = ''

    # 合并实验名称单元格
    table.rows[1].cells[1].merge(table.rows[1].cells[4])
    # 合并日期单元格
    table.rows[1].cells[5].merge(table.rows[1].cells[7])

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(11)

    doc.add_paragraph()

    # ============ 实验标题 ============
    add_formatted_paragraph(doc, '实验5  Java输入输出流与多线程编程',
                            font_name='黑体', font_size=16, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # ============ 一、实验目的 ============
    add_formatted_paragraph(doc, '一、实验目的', font_name='黑体', font_size=14, bold=True)
    purposes = [
        '1. 掌握 Java 中输入/输出流（I/O Stream）的基本概念，理解字节流和字符流的区别与适用场景。',
        '2. 掌握 FileReader/FileWriter、BufferedReader/BufferedWriter 等文件读写类的使用方法。',
        '3. 掌握 Scanner 类配合正则表达式进行文本解析和统计分析的技巧。',
        '4. 掌握 Java 多线程编程的两种方式：继承 Thread 类和实现 Runnable 接口。',
        '5. 理解线程同步机制（synchronized、wait/notify/notifyAll），解决并发访问共享数据的问题。',
        '6. 通过五个循序渐进的实验，综合运用 I/O 流和多线程知识解决实际问题。',
    ]
    for p in purposes:
        add_formatted_paragraph(doc, p, font_name='宋体', font_size=12)

    doc.add_paragraph()

    # ============ 二、实验环境 ============
    add_formatted_paragraph(doc, '二、实验环境', font_name='黑体', font_size=14, bold=True)
    env_items = [
        '操作系统：Windows 11 Home China 10.0.26200',
        '开发工具：Visual Studio Code',
        'JDK 版本：Java Development Kit (JDK)',
        '编译命令：javac -encoding UTF-8 JAVALAB/lab5/*.java',
        '运行方式：java JAVALAB.lab5.<类名>',
        '测试数据：score.txt、english.txt、hello.txt',
    ]
    for item in env_items:
        add_formatted_paragraph(doc, f'• {item}', font_name='宋体', font_size=12)
    doc.add_paragraph()

    # ============ 三、实验内容 ============
    add_formatted_paragraph(doc, '三、实验内容', font_name='黑体', font_size=14, bold=True)

    # ==================== 实验一 ====================
    add_formatted_paragraph(doc, '3.1 实验一：举重成绩单 —— 文件读写与成绩分析',
                            font_name='黑体', font_size=13, bold=True)
    add_formatted_paragraph(doc, '3.1.1 功能概述', font_name='黑体', font_size=12, bold=True)
    add_formatted_paragraph(doc,
        '本实验实现从 score.txt 读取运动员成绩，利用正则表达式和 StringTokenizer 提取数字计算总成绩，'
        '并将结果写入 scoreAnalysis.txt。综合运用 FileReader、FileWriter、BufferedReader、BufferedWriter 等字符流类。',
        font_name='宋体', font_size=12)

    add_formatted_paragraph(doc, '3.1.2 核心代码', font_name='黑体', font_size=12, bold=True)

    add_formatted_paragraph(doc, 'Fenxi.java', font_name='宋体', font_size=11, bold=True)
    with open(os.path.join(LAB5_DIR, 'Fenxi.java'), 'r', encoding='utf-8') as f:
        add_code_block(doc, f.read())

    add_formatted_paragraph(doc, 'AnalysisResult.java', font_name='宋体', font_size=11, bold=True)
    with open(os.path.join(LAB5_DIR, 'AnalysisResult.java'), 'r', encoding='utf-8') as f:
        add_code_block(doc, f.read())

    add_formatted_paragraph(doc, '3.1.3 程序运行截图', font_name='黑体', font_size=12, bold=True)
    output1, ss1 = run_and_screenshot('AnalysisResult', 'exp1', '实验1: 举重成绩单 - AnalysisResult')
    add_screenshot(doc, ss1, '图1-1 实验一运行结果')

    # 同时展示 scoreAnalysis.txt 的内容
    analysis_file = os.path.join(LAB5_DIR, 'scoreAnalysis.txt')
    if os.path.exists(analysis_file):
        with open(analysis_file, 'r', encoding='utf-8') as f:
            add_code_block(doc, f.read(), title='输出文件 scoreAnalysis.txt 内容')

    doc.add_page_break()

    # ==================== 实验二 ====================
    add_formatted_paragraph(doc, '3.2 实验二：统计英文单词 —— Scanner与正则表达式',
                            font_name='黑体', font_size=13, bold=True)
    add_formatted_paragraph(doc, '3.2.1 功能概述', font_name='黑体', font_size=12, bold=True)
    add_formatted_paragraph(doc,
        '使用 Scanner 配合正则表达式分隔符从英文文本中读取单词，利用 Vector 存储并统计单词总数、'
        '不同单词数量以及每个单词的出现频率（降序排列）。分别统计 english.txt 和 hello.txt。',
        font_name='宋体', font_size=12)

    add_formatted_paragraph(doc, '3.2.2 核心代码', font_name='黑体', font_size=12, bold=True)

    add_formatted_paragraph(doc, 'WordStatistic.java', font_name='宋体', font_size=11, bold=True)
    with open(os.path.join(LAB5_DIR, 'WordStatistic.java'), 'r', encoding='utf-8') as f:
        add_code_block(doc, f.read())

    add_formatted_paragraph(doc, 'OutputWordMess.java', font_name='宋体', font_size=11, bold=True)
    with open(os.path.join(LAB5_DIR, 'OutputWordMess.java'), 'r', encoding='utf-8') as f:
        add_code_block(doc, f.read())

    add_formatted_paragraph(doc, '3.2.3 程序运行截图', font_name='黑体', font_size=12, bold=True)
    output2, ss2 = run_and_screenshot('OutputWordMess', 'exp2', '实验2: 统计英文单词 - OutputWordMess')
    add_screenshot(doc, ss2, '图2-1 实验二运行结果')

    doc.add_page_break()

    # ==================== 实验三 ====================
    add_formatted_paragraph(doc, '3.3 实验三：密码流 —— Console密码验证与文件读取',
                            font_name='黑体', font_size=13, bold=True)
    add_formatted_paragraph(doc, '3.3.1 功能概述', font_name='黑体', font_size=12, bold=True)
    add_formatted_paragraph(doc,
        '实现控制台密码验证功能。用户最多3次机会输入密码（正确密码 tiger123），验证通过后读取 score.txt 并显示。'
        '使用 System.console().readPassword() 实现密码不回显，同时提供 IDE 环境的 Scanner 回退方案。',
        font_name='宋体', font_size=12)

    add_formatted_paragraph(doc, '3.3.2 核心代码 (PassWord.java)', font_name='黑体', font_size=12, bold=True)
    with open(os.path.join(LAB5_DIR, 'PassWord.java'), 'r', encoding='utf-8') as f:
        add_code_block(doc, f.read())

    add_formatted_paragraph(doc, '3.3.3 程序运行截图', font_name='黑体', font_size=12, bold=True)
    output3, ss3 = run_and_screenshot('PassWord', 'exp3', '实验3: 密码验证 - PassWord', input_text='tiger123\n')
    add_screenshot(doc, ss3, '图3-1 实验三运行结果——正确密码验证通过')

    doc.add_page_break()

    # ==================== 实验四 ====================
    add_formatted_paragraph(doc, '3.4 实验四：汉字输入练习 —— 双线程协作',
                            font_name='黑体', font_size=13, bold=True)
    add_formatted_paragraph(doc, '3.4.1 功能概述', font_name='黑体', font_size=12, bold=True)
    add_formatted_paragraph(doc,
        '实现双线程协作的汉字输入练习程序。GiveChineseThread（继承 Thread）每隔6秒显示一个汉字，'
        '从"好"字开始循环100个；InputChineseThread（继承 Thread）监听用户输入并与显示汉字对比打分。'
        '两个线程通过共享 Chinese 对象传递数据。输入 # 结束程序。',
        font_name='宋体', font_size=12)

    add_formatted_paragraph(doc, '3.4.2 类设计', font_name='黑体', font_size=12, bold=True)
    add_formatted_paragraph(doc,
        '• Chinese 类：共享数据类，封装一个 char 变量，提供 setChinese/getChinese 方法\n'
        '• GiveChineseThread 类：生产者线程，循环更新 Chinese 对象中的汉字\n'
        '• InputChineseThread 类：消费者线程，读取用户输入并与 Chinese 对象比对\n'
        '• TypeChinese 类：主程序，创建共享对象和两个线程并启动',
        font_name='宋体', font_size=12)

    add_formatted_paragraph(doc, '3.4.3 核心代码', font_name='黑体', font_size=12, bold=True)
    for fname, label in [
        ('TypeChinese.java', 'TypeChinese.java（主程序）'),
        ('Chinese.java', 'Chinese.java（共享数据类）'),
        ('GiveChineseThread.java', 'GiveChineseThread.java（出汉字线程）'),
        ('InputChineseThread.java', 'InputChineseThread.java（输入线程）'),
    ]:
        add_formatted_paragraph(doc, label, font_name='宋体', font_size=11, bold=True)
        with open(os.path.join(LAB5_DIR, fname), 'r', encoding='utf-8') as f:
            add_code_block(doc, f.read())

    add_formatted_paragraph(doc, '3.4.4 程序运行截图', font_name='黑体', font_size=12, bold=True)
    output4, ss4 = run_and_screenshot('TypeChinese', 'exp4', '实验4: 汉字输入练习 - TypeChinese',
                                       input_text='好\n好\n坏\n#\n', timeout=5)
    add_screenshot(doc, ss4, '图4-1 实验四运行结果——输入"好"得2分，输入"坏"未得分，输入#退出')

    doc.add_page_break()

    # ==================== 实验五 ====================
    add_formatted_paragraph(doc, '3.5 实验五：双线程猜数字 —— 线程同步与wait/notify通信',
                            font_name='黑体', font_size=13, bold=True)
    add_formatted_paragraph(doc, '3.5.1 功能概述', font_name='黑体', font_size=12, bold=True)
    add_formatted_paragraph(doc,
        '实现双线程猜数字游戏。Number 类实现 Runnable 接口，giveNumberThread 生成1-100随机数并反馈"猜大/猜小"，'
        'guessNumberThread 使用二分查找逐步缩小范围。两个线程通过 synchronized 方法和 wait/notifyAll 机制协调通信：'
        '出题线程等待猜数结果，猜数线程等待反馈信号。这是 Java 并发编程的经典案例。',
        font_name='宋体', font_size=12)

    add_formatted_paragraph(doc, '3.5.2 关键技术点', font_name='黑体', font_size=12, bold=True)
    key_points = [
        '• 实现 Runnable 接口创建线程，使多个线程共享同一个 Number 实例',
        '• synchronized 同步方法 setMessage() 确保互斥访问共享数据',
        '• wait() 释放锁并进入等待状态；notifyAll() 唤醒所有等待线程',
        '• pleaseGuess 布尔变量作为线程间"信号灯"，控制执行顺序',
        '• 二分查找算法：维护 min/max 范围折半，O(log n) 效率（最多7次）',
    ]
    for kp in key_points:
        add_formatted_paragraph(doc, kp, font_name='宋体', font_size=12)

    add_formatted_paragraph(doc, '3.5.3 核心代码', font_name='黑体', font_size=12, bold=True)
    add_formatted_paragraph(doc, 'TwoThreadGuessNumber.java（主程序）', font_name='宋体', font_size=11, bold=True)
    with open(os.path.join(LAB5_DIR, 'TwoThreadGuessNumber.java'), 'r', encoding='utf-8') as f:
        add_code_block(doc, f.read())

    add_formatted_paragraph(doc, 'Number.java（核心——线程同步与通信）', font_name='宋体', font_size=11, bold=True)
    with open(os.path.join(LAB5_DIR, 'Number.java'), 'r', encoding='utf-8') as f:
        add_code_block(doc, f.read())

    add_formatted_paragraph(doc, '3.5.4 程序运行截图', font_name='黑体', font_size=12, bold=True)
    output5, ss5 = run_and_screenshot('TwoThreadGuessNumber', 'exp5', '实验5: 双线程猜数字 - TwoThreadGuessNumber')
    add_screenshot(doc, ss5, '图5-1 实验五运行结果——二分法约6次猜对')

    # ==================== 输入数据文件 ====================
    doc.add_page_break()
    add_formatted_paragraph(doc, '3.6 测试数据文件', font_name='黑体', font_size=13, bold=True)
    for fname in ['score.txt', 'english.txt', 'hello.txt']:
        add_formatted_paragraph(doc, fname, font_name='黑体', font_size=12, bold=True)
        with open(os.path.join(LAB5_DIR, fname), 'r', encoding='utf-8') as f:
            add_code_block(doc, f.read())

    # ==================== 四、心得体会 ====================
    doc.add_page_break()
    add_formatted_paragraph(doc, '四、心得体会', font_name='黑体', font_size=14, bold=True)

    thoughts = [
        ('（一）Java I/O 流方面',
         '通过实验一至实验三，我对 Java 输入输出流有了系统的认识和实践：\n\n'
         '1. 理解了字节流与字符流的区别：字节流以字节为单位处理二进制数据，而字符流以字符为单位处理文本数据。'
         '实验中使用的 FileReader/FileWriter 是字符流的便捷实现，自动处理了字符编码问题。\n\n'
         '2. 掌握了缓冲流的装饰器模式：BufferedReader 和 BufferedWriter 通过包装基础的 Reader/Writer，'
         '引入缓冲区显著减少实际 I/O 操作次数。readLine() 和 newLine() 方法使逐行读写非常便捷。\n\n'
         '3. 学会了 Scanner 配合正则表达式的文本解析：useDelimiter(regex) 方法可灵活设置分隔符，'
         '正则 [\\s\\d\\p{Punct}]+ 能一次性匹配空格、数字和标点符号，是处理自然语言文本的常用模式。\n\n'
         '4. StringTokenizer 与 replaceAll 的互补性：先用正则去除非数字字符，再用 StringTokenizer 拆分，'
         '是解析混合文本中数值信息的有效方法。'),

        ('（二）Java 多线程编程方面',
         '通过实验四和实验五，我对 Java 多线程编程有了深入的理解和实践：\n\n'
         '1. 两种多线程实现方式对比：继承 Thread 类简单直接但受限于单继承；实现 Runnable 接口更灵活且适合资源共享。'
         '实验四中两个线程类继承 Thread，实验五中 Number 类实现 Runnable 让两个线程共享同一实例。\n\n'
         '2. 线程同步机制的核心作用：当多线程并发访问共享数据时，synchronized 关键字确保操作原子性。'
         '实验五的 setMessage() 方法被 synchronized 修饰，避免数据不一致。\n\n'
         '3. wait/notify 线程通信：wait() 释放锁并进入等待，notifyAll() 唤醒等待线程。'
         '配合 pleaseGuess 布尔标志位实现精确的"生产者-消费者"交替执行模型。\n\n'
         '4. 二分查找算法的应用：猜数线程每次取中间值，将1~100的猜测次数从最坏100次降到最多7次（log₂100≈6.64）。'),

        ('（三）综合收获',
         '本次实验五个任务总计约400行代码，循序渐进地覆盖了 Java I/O 流（FileReader/FileWriter、'
         'BufferedReader/BufferedWriter、Scanner、Console）、正则表达式、StringTokenizer、'
         '多线程编程（Thread、Runnable、synchronized、wait/notify）、集合框架（Vector）等核心知识点。\n\n'
         '实验设计由浅入深：文件读写→文本解析统计→密码验证→双线程协作→线程同步通信，'
         '帮助我建立了对 Java I/O 和并发编程的系统认知。\n\n'
         '调试中也遇到典型问题：IDE 中 System.console() 返回 null、Windows 终端 UTF-8 编码显示问题、'
         '多线程程序执行顺序的不确定性等，促使我深入理解 Java 的跨平台特性和并发编程复杂性。'),
    ]

    for subject, content in thoughts:
        add_formatted_paragraph(doc, subject, font_name='黑体', font_size=13, bold=True)
        for line in content.split('\n'):
            if line.strip():
                add_formatted_paragraph(doc, line.strip(), font_name='宋体', font_size=12)
        doc.add_paragraph()

    # ==================== 附录 ====================
    doc.add_paragraph()
    add_formatted_paragraph(doc, '附录：项目文件清单', font_name='黑体', font_size=14, bold=True)

    file_list = [
        ('实验一', 'AnalysisResult.java', '举重成绩单主程序', '~30行'),
        ('实验一', 'Fenxi.java', '分数分析工具类', '~20行'),
        ('实验二', 'WordStatistic.java', '单词统计类', '~50行'),
        ('实验二', 'OutputWordMess.java', '单词统计主程序', '~55行'),
        ('实验三', 'PassWord.java', '密码验证与文件读取', '~50行'),
        ('实验四', 'TypeChinese.java', '汉字练习主程序', '~30行'),
        ('实验四', 'Chinese.java', '共享汉字数据类', '~15行'),
        ('实验四', 'GiveChineseThread.java', '出汉字线程', '~30行'),
        ('实验四', 'InputChineseThread.java', '输入汉字线程', '~30行'),
        ('实验五', 'TwoThreadGuessNumber.java', '双线程猜数字主程序', '~10行'),
        ('实验五', 'Number.java', '猜数字核心类(同步)', '~70行'),
        ('——', 'score.txt / english.txt / hello.txt', '测试数据文件', '共7行'),
    ]

    ftable = doc.add_table(rows=len(file_list)+1, cols=4, style='Table Grid')
    ftable.autofit = True
    for i, h in enumerate(['所属实验', '文件名', '说明', '代码行数']):
        ftable.rows[0].cells[i].text = h
        for p in ftable.rows[0].cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.size = Pt(10)

    for row_idx, (exp, fname, desc, lines) in enumerate(file_list):
        for col_idx, val in enumerate([exp, fname, desc, lines]):
            ftable.rows[row_idx+1].cells[col_idx].text = val

    for row in ftable.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(10)

    # ==================== 保存 ====================
    output_path = os.path.join(LAB5_DIR, '实验5_Java输入输出流与多线程编程_软件工程3班_3242052051355_洪耀东.docx')
    doc.save(output_path)
    print(f"Word report generated: {output_path}")
    print(f"Screenshots saved to: {SCREENSHOTS_DIR}")

if __name__ == '__main__':
    main()
