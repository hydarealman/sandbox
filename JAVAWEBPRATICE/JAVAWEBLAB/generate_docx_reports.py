#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate all 10 JavaWeb experiment reports as .docx Word documents."""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = r'd:\sandbox\JAVAWEBPRATICE\JAVAWEBLAB'
REPORTS_DIR = os.path.join(BASE_DIR, 'reports')
SCREENSHOTS_DIR = os.path.join(REPORTS_DIR, 'screenshots')

EXPERIMENTS = [
    {
        'num': 1,
        'title': '调查问卷制作',
        'subtitle': 'HTML常用标签的应用',
        'purpose': [
            '1. 掌握HTML的常用标签；',
            '2. 认真书写实验报告，如实填写各项实验内容。'
        ],
        'content': [
            '编写网页practice1_1.html，具体要求如下：',
            '（1）网页中有个form表单，表单处理程序为本页面程序，表单提交方式为"post"，表单提交编码方式为"multipart/form-data"；',
            '（2）页面包含用户名、密码、性别（单选）、年龄（下拉选择）、爱好（多选）、学历（下拉选择）、自我介绍（文本域）等表单元素；',
            '（3）使用CSS对页面进行美化，包含提交和重置按钮。'
        ],
        'source_files': [
            ('experiment01/practice1_1.html', 'practice1_1.html - 调查问卷页面')
        ],
        'screenshots': ['experiment01.png'],
        'notes': '本次实验通过编写HTML表单页面，掌握了form标签、input标签（text/password/radio/checkbox）、select标签、textarea标签等常用HTML标签的使用方法。同时学习了CSS基本样式设置，包括字体、颜色、边框、背景等属性的应用。通过实践加深了对HTML表单元素和CSS样式的理解，为后续JSP动态网页开发打下了基础。',
        'deploy': '直接使用浏览器打开 experiment01/practice1_1.html 文件即可运行'
    },
    {
        'num': 2,
        'title': '页面布局',
        'subtitle': 'DIV+CSS布局类型',
        'purpose': [
            '1. 掌握CSS的基本语法；',
            '2. 掌握常见的DIV+CSS布局类型；',
            '3. 认真书写实验报告，如实填写各项实验内容。'
        ],
        'content': [
            '按照以下步骤编写网页practice1_2.html：',
            '步骤1：使用DIV定义结构——页头(#header)、导航栏(#navigator)、内容(#content)、版权(#footer)，全部放入body容器中（width:800px）；',
            '步骤2：定义body属性——width:800px, margin:0px auto居中, border边框；',
            '步骤3：定义页头#header——height:100px，背景图片background-image；',
            '步骤4：定义导航栏#navigator——横向排列6个链接（首页/文章/相册/Blog/论坛/帮助），float:left布局，hover变色效果；',
            '步骤5：定义内容#content——文本缩进text-indent:2em，行高line-height:1.5em，文章段落排版；',
            '步骤6：定义页脚#footer——居中版权信息，背景色#009966。'
        ],
        'source_files': [
            ('experiment02/practice1_2.html', 'practice1_2.html - DIV+CSS页面布局')
        ],
        'screenshots': ['experiment02.png'],
        'notes': '本次实验学习了DIV+CSS进行网页布局的方法。通过将页面划分为header、navigator、content、footer四个结构块，使用CSS的float属性实现导航栏的横向排列，使用margin:0px auto实现整体居中。理解了CSS盒子模型的概念以及CSS选择器、伪类（:hover）的使用方法。',
        'deploy': '直接使用浏览器打开 experiment02/practice1_2.html 文件即可运行'
    },
    {
        'num': 3,
        'title': '表单验证',
        'subtitle': 'JavaScript事件处理机制',
        'purpose': [
            '1. 掌握JavaScript的语法基础；',
            '2. 掌握JavaScript的事件处理机制；',
            '3. 认真书写实验报告，如实填写各项实验内容。'
        ],
        'content': [
            '制作一个用户注册页面practice1_3.html，具体要求如下：',
            '（1）包含登录账号、密码、确认密码、姓名、身份证号码（18位）、出生年月日、住址、邮编、E-mail等输入区域；',
            '（2）在提交时使用JavaScript（onsubmit事件）检验输入是否符合要求，alert出非法的输入，并将焦点返回对应控件；',
            '（3）根据出生年月日判断身份证号码是否合法（只考虑18位身份证，提取第7-14位与出生日期比对）；',
            '（4）验证规则：登录账号只能是字母或数字且以字母开头；密码8位以上且需要有字母、数字和特殊字符；出生年月日格式为yyyy-mm-dd；邮编为6位数字；E-mail基本格式验证（包含@和.）。'
        ],
        'source_files': [
            ('experiment03/practice1_3.html', 'practice1_3.html - JavaScript表单验证')
        ],
        'screenshots': ['experiment03.png'],
        'notes': '本次实验学习了JavaScript的表单验证技术。掌握了onsubmit事件处理、正则表达式验证、字符串操作、DOM元素访问（getElementById）等知识点。特别是身份证号码与出生日期的交叉验证，加深了对JavaScript逻辑编程和数据处理的理解。表单验证是Web开发中的重要环节，有效提高了用户体验和数据质量。',
        'deploy': '直接使用浏览器打开 experiment03/practice1_3.html 文件即可运行'
    },
    {
        'num': 4,
        'title': 'JSP脚本元素的应用',
        'subtitle': 'Java程序片、成员变量与表达式',
        'purpose': [
            '1. 掌握Java程序片的定义与使用（<% %>）；',
            '2. 掌握JSP成员变量与方法的定义与使用（<%! %>）；',
            '3. 掌握Java表达式的定义与使用（<%= %>）；',
            '4. 认真书写实验报告，如实填写各项实验内容。'
        ],
        'content': [
            '编写1个JSP页面index.jsp，在JSP页面中使用以下JSP脚本元素：',
            '1. 使用JSP表达式(<%= %>)输出当前日期时间和系统信息；',
            '2. 使用Java程序片（<% %>）实现循环输出，在浏览器中输出大小为15*10的表格，以及26个小写英文字母表；',
            '3. 利用成员变量（<%! %>）被所有客户共享这一性质，实现一个简单的页面访问计数器；',
            '4. 使用JSP声明（<%! %>）定义getGreeting()方法，根据时间段返回不同的问候语；',
            '5. 综合演示JSP三种脚本元素（声明、脚本段、表达式）的使用。'
        ],
        'source_files': [
            ('experiment04/index.jsp', 'index.jsp - JSP脚本元素综合演示')
        ],
        'screenshots': ['experiment04.png'],
        'notes': '本次实验深入学习了JSP的三种脚本元素：声明（<%! %>）用于定义成员变量和方法，属于Servlet的实例变量，被所有请求共享；脚本段（<% %>）用于编写Java逻辑代码，位于_jspService()方法中；表达式（<%= %>）用于输出数据，等价于out.print()。通过访问计数器的实现，深刻理解了JSP成员变量的共享特性。',
        'deploy': '部署在Tomcat webapps/experiment04/目录下，通过 http://localhost:8088/experiment04/ 访问'
    },
    {
        'num': 5,
        'title': 'JSP动作标记的应用',
        'subtitle': 'include、forward和param动作标记',
        'purpose': [
            '1. 掌握动作标记include、forward和param的应用；',
            '2. 认真书写实验报告，如实填写各项实验内容。'
        ],
        'content': [
            '1. include和param动作标记的应用：编写3个JSP页面input.jsp、first.jsp和second.jsp。input.jsp使用<jsp:include>动作标记加载first.jsp和second.jsp页面，并通过<jsp:param>传递参数。first.jsp根据param提供的行数和列数生成动态表格；second.jsp根据param提供的两个正整数计算最大公约数（辗转相除法）和最小公倍数。',
            '2. 登录验证：编写3个JSP页面login.jsp、validate.jsp和success.jsp。login.jsp输入用户名与密码，提交给validate.jsp进行验证（合法用户：tom，密码：jenny），如果验证通过则使用<jsp:forward>转发到success.jsp显示欢迎信息，否则forward回login.jsp并显示错误提示。'
        ],
        'source_files': [
            ('experiment05/input.jsp', 'input.jsp - include和param动作标记加载页面'),
            ('experiment05/first.jsp', 'first.jsp - 动态生成表格'),
            ('experiment05/second.jsp', 'second.jsp - 计算最大公约数和最小公倍数'),
            ('experiment05/login.jsp', 'login.jsp - 登录页面'),
            ('experiment05/validate.jsp', 'validate.jsp - 验证页面（使用forward动作标记）'),
            ('experiment05/success.jsp', 'success.jsp - 登录成功页面')
        ],
        'screenshots': ['experiment05_input.png', 'experiment05_login.png'],
        'notes': '本次实验学习了JSP动作标记的使用。通过<jsp:include>实现了页面的模块化包含，通过<jsp:param>实现了参数传递，使得被包含的页面能够动态获取参数并生成不同的输出内容。同时学习了<jsp:forward>的使用，理解了服务端转发（forward）和客户端跳转（redirect）的区别：forward在服务器内部完成，URL不变，可以共享request；redirect是客户端重新发起请求，URL改变。',
        'deploy': '部署在Tomcat webapps/experiment05/目录下，通过 http://localhost:8088/experiment05/input.jsp 和 http://localhost:8088/experiment05/login.jsp 访问'
    },
    {
        'num': 6,
        'title': 'JSP内置对象的应用',
        'subtitle': 'request、response、session、application对象',
        'purpose': [
            '1. 掌握JSP内置对象的定义及常用方法；',
            '2. 理解request、session以及application的区别；',
            '3. 认真书写实验报告，如实填写各项实验内容。'
        ],
        'content': [
            '1. 编写index.jsp页面，综合演示JSP四大内置对象（request、response、session、application）的常用方法。包括：request获取请求参数和客户端信息（getParameter、getRemoteAddr、getHeader等）；response设置响应头；session属性存取和会话跟踪（getAttribute/setAttribute/getId）；application全局属性存取和应用级计数器。并展示四种作用域的比较表格。',
            '2. 编写shop.jsp页面，实现基于HttpSession的购物车功能。商品包括：《Java编程思想》、《JSP实战指南》、《Servlet核心技术》、《Web开发大全》。用户可以添加商品到购物车（数量累加）、查看购物车列表、清空购物车。购物车采用Map<String,Integer>结构存储在session中。'
        ],
        'source_files': [
            ('experiment06/index.jsp', 'index.jsp - JSP内置对象综合演示'),
            ('experiment06/shop.jsp', 'shop.jsp - 基于Session的购物车')
        ],
        'screenshots': ['experiment06_index.png', 'experiment06_shop.png'],
        'notes': '本次实验深入理解了JSP四大内置对象的区别：request的作用域仅限于一次请求（请求结束后销毁）；response用于向客户端发送响应；session的作用域是一次会话（同一用户在多个请求间共享数据）；application的作用域是整个Web应用（所有用户共享）。通过购物车的实现，加深了对Session跟踪机制的理解，掌握了会话状态管理的基本方法。',
        'deploy': '部署在Tomcat webapps/experiment06/目录下，通过 http://localhost:8088/experiment06/ 访问'
    },
    {
        'num': 7,
        'title': 'JavaBean的应用',
        'subtitle': 'jsp:useBean、jsp:setProperty、jsp:getProperty',
        'purpose': [
            '1. 掌握JavaBean的定义规范（无参构造、私有属性、getter/setter方法）；',
            '2. 掌握JSP中useBean、setProperty、getProperty动作标记的使用；',
            '3. 认真书写实验报告，如实填写各项实验内容。'
        ],
        'content': [
            '编写三角形JavaBean（Triangle.java）和JSP页面（inputTriangle.jsp、showTriangle.jsp），实现三角形计算器。',
            'Triangle.java包含三个double属性：sideA、sideB、sideC。提供方法：isValid()判断三角形是否合法（任意两边之和大于第三边），getPerimeter()计算周长，getArea()使用海伦公式（s=半周长，面积=√(s(s-a)(s-b)(s-c))）计算面积。',
            'inputTriangle.jsp提供表单输入三边长，提交到showTriangle.jsp。showTriangle.jsp使用<jsp:useBean id="triangle" class="beans.Triangle" scope="request">实例化Bean，<jsp:setProperty name="triangle" property="*">自动将请求参数填充到Bean属性中，<jsp:getProperty>显示各边长，调用方法显示周长和面积。'
        ],
        'source_files': [
            ('experiment07/WEB-INF/classes/Triangle.java', 'Triangle.java - 三角形JavaBean'),
            ('experiment07/inputTriangle.jsp', 'inputTriangle.jsp - 输入三角形边长'),
            ('experiment07/showTriangle.jsp', 'showTriangle.jsp - 显示三角形计算结果')
        ],
        'screenshots': ['experiment07_input.png', 'experiment07_show.png'],
        'notes': '本次实验学习了JavaBean的编写规范和使用方法。JavaBean需要有无参构造方法、私有属性、对应的getter/setter方法。通过<jsp:useBean>创建或获取Bean实例，<jsp:setProperty property="*">可以自动将请求参数名与Bean属性名匹配并填充，大大简化了参数获取代码。海伦公式的实现加深了对Java面向对象编程的理解，同时也练习了JSP动作标记与JavaBean的集成使用。',
        'deploy': '部署在Tomcat webapps/experiment07/目录下，Triangle.java编译为WEB-INF/classes/beans/Triangle.class，通过 http://localhost:8088/experiment07/inputTriangle.jsp 访问'
    },
    {
        'num': 8,
        'title': 'Servlet的应用',
        'subtitle': 'Servlet登录验证、forward与redirect区别',
        'purpose': [
            '1. 掌握Servlet的编写与配置（@WebServlet注解）；',
            '2. 理解RequestDispatcher.forward()与response.sendRedirect()的区别；',
            '3. 认真书写实验报告，如实填写各项实验内容。'
        ],
        'content': [
            '编写一个Servlet（LoginServlet_1.java）实现登录验证，使用@WebServlet注解配置URL映射为/loginServlet_1。',
            'login_1.jsp提供登录表单（method="post" action="loginServlet_1"），提交到LoginServlet_1处理。',
            'LoginServlet_1在doPost()方法中：先验证用户名和密码是否为空；然后验证是否为合法用户（用户名：zhangsan，密码：123）。如果登录成功，使用RequestDispatcher.forward()将请求转发到loginSuccess_1.jsp，并设置request属性传递用户名和密码；如果登录失败，使用response.sendRedirect()重定向回login_1.jsp并附带error参数。',
            'loginSuccess_1.jsp显示登录成功信息，展示用户名和密码，并在页面底部说明forward与redirect的区别（forward：服务器内部转发，URL不变，可以传递request属性；redirect：客户端重定向，URL改变，通过URL参数传递信息）。'
        ],
        'source_files': [
            ('experiment08/WEB-INF/classes/LoginServlet_1.java', 'LoginServlet_1.java - Servlet登录验证'),
            ('experiment08/login_1.jsp', 'login_1.jsp - 登录表单页面'),
            ('experiment08/loginSuccess_1.jsp', 'loginSuccess_1.jsp - 登录成功页面')
        ],
        'screenshots': ['experiment08_login.png'],
        'notes': '本次实验学习了Servlet的编写方法，掌握了@WebServlet注解配置URL映射的方式（无需web.xml）。通过实际编码深入理解了forward和redirect的核心区别：①forward是服务器内部行为，客户端感知不到，浏览器地址栏URL不变，可以在request域中共享数据；②redirect是服务器通知客户端重新发起请求，浏览器地址栏URL变为新地址，不能直接传递request属性，需要通过URL参数或session传递数据。',
        'deploy': '部署在Tomcat webapps/experiment08/目录下，LoginServlet_1.java编译为WEB-INF/classes/LoginServlet_1.class，通过 http://localhost:8088/experiment08/login_1.jsp 访问'
    },
    {
        'num': 9,
        'title': 'MVC模式的应用',
        'subtitle': 'Model（JavaBean）+ View（JSP）+ Controller（Servlet）',
        'purpose': [
            '1. 理解MVC设计模式的基本思想和各层职责；',
            '2. 掌握使用Servlet+JSP+JavaBean实现MVC模式的方法；',
            '3. 认真书写实验报告，如实填写各项实验内容。'
        ],
        'content': [
            '使用MVC模式实现一个简单的计算器应用，分为三层：',
            'Model层（模型）：Computer.java为实体模型，封装两个操作数（operand1、operand2）、运算符（operator）和计算结果（result）；CalculateBean.java为业务模型，包含calculate(Computer)方法，根据运算符执行加减乘除运算（除法需检查除数不为0）。',
            'View层（视图）：inputNumber.jsp为输入页面，提供两个数字输入框和一个运算符下拉选择框（+、-、*、/），表单提交到handleComputer；showResult.jsp为结果展示页面，显示算式和结果，或显示错误信息。',
            'Controller层（控制器）：HandleComputer.java为Servlet（@WebServlet("/handleComputer")），在doPost()中获取请求参数，创建Computer实体并设置属性，调用CalculateBean.calculate()执行计算，将结果存入request属性，forward到showResult.jsp显示。'
        ],
        'source_files': [
            ('experiment09/WEB-INF/classes/model/Computer.java', 'Computer.java - 实体模型（Model）'),
            ('experiment09/WEB-INF/classes/model/CalculateBean.java', 'CalculateBean.java - 业务模型（Model）'),
            ('experiment09/WEB-INF/classes/model/HandleComputer.java', 'HandleComputer.java - Servlet控制器（Controller）'),
            ('experiment09/inputNumber.jsp', 'inputNumber.jsp - 计算器输入页面（View）'),
            ('experiment09/showResult.jsp', 'showResult.jsp - 计算结果页面（View）')
        ],
        'screenshots': ['experiment09_input.png'],
        'notes': '本次实验完整实践了MVC设计模式。Model层（Computer+CalculateBean）负责数据封装和业务逻辑处理，与界面无关；View层（JSP页面）专注于页面展示；Controller层（HandleComputer Servlet）作为协调者，接收请求、调用业务逻辑、选择视图。这种分层架构使代码职责清晰、耦合度低、便于维护和扩展，是JavaWeb开发中的核心设计模式。',
        'deploy': '部署在Tomcat webapps/experiment09/目录下，Java类编译到WEB-INF/classes/model/，通过 http://localhost:8088/experiment09/inputNumber.jsp 访问'
    },
    {
        'num': 10,
        'title': '会话跟踪技术',
        'subtitle': 'HttpSession、URL重写、session.invalidate()注销',
        'purpose': [
            '1. 掌握HttpSession会话跟踪技术的使用方法；',
            '2. 掌握URL重写（response.encodeURL()/encodeRedirectURL()）的使用场景和原理；',
            '3. 理解session.invalidate()注销机制；',
            '4. 认真书写实验报告，如实填写各项实验内容。'
        ],
        'content': [
            '实现一个基于HttpSession的完整登录注销系统，包含以下组件：',
            '1. LoginServlet.java（@WebServlet("/login")）：处理登录POST请求，验证用户（规则：用户名与密码相同即为合法用户，如admin/admin），创建HttpSession并保存用户信息（username、loginTime、sessionId），使用response.encodeRedirectURL()进行URL重写后重定向到welcome.jsp；',
            '2. LogoutServlet.java（@WebServlet("/logout")）：处理注销GET请求，通过request.getSession(false)获取当前会话（不创建新会话），调用session.invalidate()销毁会话数据，重定向回login.jsp并显示"已成功注销"提示；',
            '3. login.jsp：登录表单页面（action="login" method="post"），接收并显示error（红色）和info（蓝色）参数信息；',
            '4. welcome.jsp：欢迎页面，首先检查session中是否有username属性（无则重定向到登录页），显示用户欢迎信息和完整的session详情（session ID、登录时间、创建时间、最后访问时间、最大不活动间隔），所有内部链接使用response.encodeURL()进行URL重写，确保在浏览器禁用Cookie时仍能正常使用；',
            '5. profile.jsp：个人信息页面，同样需要登录验证，展示用户资料表格（用户名、登录时间、session ID、新建状态、编码URL），所有导航链接使用encodeURL处理。'
        ],
        'source_files': [
            ('experiment10/WEB-INF/classes/LoginServlet.java', 'LoginServlet.java - 登录Servlet'),
            ('experiment10/WEB-INF/classes/LogoutServlet.java', 'LogoutServlet.java - 注销Servlet'),
            ('experiment10/login.jsp', 'login.jsp - 登录页面'),
            ('experiment10/welcome.jsp', 'welcome.jsp - 欢迎页面（含URL重写）'),
            ('experiment10/profile.jsp', 'profile.jsp - 个人信息页面（含URL重写）')
        ],
        'screenshots': ['experiment10_login.png', 'experiment10_welcome.png', 'experiment10_profile.png'],
        'notes': '本次实验深入学习了HttpSession会话跟踪技术。通过完整的登录注销系统实现，理解了以下关键知识点：①session的创建（request.getSession(true)）和获取（request.getSession(false)）；②session属性的存取（setAttribute/getAttribute）；③session的生命周期管理（创建时间、最后访问时间、最大不活动间隔）；④session的销毁（invalidate()方法）。URL重写技术（response.encodeURL()/encodeRedirectURL()）是本次实验的另一个重点，它通过在URL后追加;jsessionid=xxx参数，确保在浏览器禁用Cookie的情况下session仍能正常工作。这是Web应用健壮性的重要保障。',
        'deploy': '部署在Tomcat webapps/experiment10/目录下，Servlet编译到WEB-INF/classes/，通过 http://localhost:8088/experiment10/login.jsp 访问'
    }
]

def set_cell_font(cell, text, font_name='宋体', font_size=Pt(11), bold=False):
    """Set cell text with formatting."""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = font_size
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold

def add_heading_styled(doc, text, level=1):
    """Add a heading with Chinese-friendly font."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return heading

def add_paragraph_styled(doc, text, font_name='宋体', font_size=Pt(12), bold=False, indent=True):
    """Add a paragraph with proper Chinese font."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)  # 2字符缩进
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.size = font_size
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold
    return p

def add_code_block(doc, code_text, filename=''):
    """Add a code block with dark background and monospace font."""
    # Add filename label
    if filename:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run('📄 ' + filename)
        run.font.size = Pt(9)
        run.font.name = 'Consolas'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.color.rgb = RGBColor(100, 100, 100)

    # Add code with shading
    for line in code_text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        # Add grey background
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F0F0F0')
        p.paragraph_format.element.get_or_add_pPr().append(shd)

        run = p.add_run(line if line else ' ')
        run.font.size = Pt(8)
        run.font.name = 'Consolas'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.color.rgb = RGBColor(30, 30, 30)

def read_source_file(rel_path):
    """Read source file content."""
    full_path = os.path.join(BASE_DIR, rel_path)
    if os.path.exists(full_path):
        with open(full_path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    return f'[文件未找到: {rel_path}]'

def add_info_table(doc):
    """Add the header info table (班级/学号/姓名/日期)."""
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'

    # Row 1
    set_cell_font(table.rows[0].cells[0], '班级', '宋体', Pt(11), True)
    set_cell_font(table.rows[0].cells[1], '', '宋体', Pt(11))
    set_cell_font(table.rows[0].cells[2], '学号', '宋体', Pt(11), True)
    set_cell_font(table.rows[0].cells[3], '', '宋体', Pt(11))

    # Row 2
    set_cell_font(table.rows[1].cells[0], '姓名', '宋体', Pt(11), True)
    set_cell_font(table.rows[1].cells[1], '', '宋体', Pt(11))
    set_cell_font(table.rows[1].cells[2], '实验日期', '宋体', Pt(11), True)
    set_cell_font(table.rows[1].cells[3], '2026年6月16日', '宋体', Pt(11))

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(2)
        row.cells[1].width = Cm(6)
        row.cells[2].width = Cm(2)
        row.cells[3].width = Cm(6)

    doc.add_paragraph()  # spacing

def generate_docx(exp):
    """Generate a single experiment report as .docx."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Style setup
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5

    # Title
    title = doc.add_heading(f'实验{exp["num"]} {exp["title"]}', level=0)
    for run in title.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f'—— {exp["subtitle"]} ——')
    run.font.size = Pt(14)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()  # spacing

    # Info table
    add_info_table(doc)

    # 一、实验目的
    add_heading_styled(doc, '一、实验目的', level=2)
    for item in exp['purpose']:
        add_paragraph_styled(doc, item)

    # 二、实验环境
    add_heading_styled(doc, '二、实验环境', level=2)
    env_items = [
        '操作系统：Windows 11 Home China',
        'JDK版本：JDK 25.0.2 (Temurin)',
        'Web服务器：Apache Tomcat 11.0.22 (Jakarta EE)',
        '开发工具：VS Code / 文本编辑器',
        '浏览器：Microsoft Edge',
        f'运行方式：{exp["deploy"]}'
    ]
    for item in env_items:
        add_paragraph_styled(doc, item, indent=False)

    # 三、实验内容
    add_heading_styled(doc, '三、实验内容', level=2)
    for item in exp['content']:
        add_paragraph_styled(doc, item)

    # 四、实验练习（源码）
    add_heading_styled(doc, '四、实验练习——程序源代码', level=2)
    for i, (file_path, desc) in enumerate(exp['source_files']):
        add_paragraph_styled(doc, f'{i+1}. {desc}', '黑体', Pt(11), True, indent=False)
        code = read_source_file(file_path)
        add_code_block(doc, code, file_path)
        doc.add_paragraph()  # spacing

    # 五、运行结果截图
    add_heading_styled(doc, '五、运行结果截图', level=2)
    for i, shot in enumerate(exp['screenshots']):
        shot_path = os.path.join(SCREENSHOTS_DIR, shot)
        if os.path.exists(shot_path):
            add_paragraph_styled(doc, f'图{i+1}：{exp["title"]} - 运行效果截图', '黑体', Pt(10), True, indent=False)
            try:
                doc.add_picture(shot_path, width=Inches(5.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                add_paragraph_styled(doc, f'[图片加载失败: {e}]', indent=False)
            doc.add_paragraph()
        else:
            add_paragraph_styled(doc, f'[截图未找到: {shot}]', indent=False)

    # 六、实验心得
    add_heading_styled(doc, '六、实验心得', level=2)
    add_paragraph_styled(doc, exp['notes'])

    # Save
    filename = f'实验{exp["num"]:02d}_{exp["title"]}.docx'
    filepath = os.path.join(REPORTS_DIR, filename)
    doc.save(filepath)
    return filepath

def main():
    os.makedirs(REPORTS_DIR, exist_ok=True)

    for exp in EXPERIMENTS:
        num = exp['num']
        print(f'Generating docx for experiment {num:02d}...')
        filepath = generate_docx(exp)
        size_kb = os.path.getsize(filepath) / 1024
        print(f'  [OK] {os.path.basename(filepath)} ({size_kb:.0f} KB)')

    print(f'\n[OK] All {len(EXPERIMENTS)} .docx reports generated!')
    print(f'  Reports directory: {REPORTS_DIR}')

if __name__ == '__main__':
    main()
